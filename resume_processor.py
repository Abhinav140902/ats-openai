import os
import re
from typing import List, Tuple, Dict
import pypdf
import docx
from rich.console import Console
from config import Config

console = Console()

class ResumeProcessor:
    def __init__(self):
        self.resume_dir = Config.RESUME_DIR
        os.makedirs(self.resume_dir, exist_ok=True)
        
    def load_resumes(self) -> List[Tuple[str, str]]:
        """Load all PDF and DOCX resumes."""
        resume_texts = []
        
        files = [f for f in os.listdir(self.resume_dir) 
                if f.lower().endswith(('.pdf', '.docx'))]
        
        if not files:
            console.print(f"[red]No PDF/DOCX files found in {self.resume_dir}[/red]")
            return resume_texts
        
        console.print(f"[green]Processing {len(files)} resume files...[/green]")
        
        for filename in files:
            file_path = os.path.join(self.resume_dir, filename)
            
            try:
                if filename.lower().endswith('.pdf'):
                    text = self._extract_from_pdf(file_path)
                else:
                    text = self._extract_from_docx(file_path)
                
                if text and text.strip():
                    resume_texts.append((filename, text))
                    console.print(f"✅ {filename}: {len(text)} chars")
            except Exception as e:
                console.print(f"[red]❌ {filename}: {str(e)}[/red]")
        
        return resume_texts
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        text = ""
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = docx.Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text).strip()
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from resume."""
        sections = {
            'skills': '',
            'experience': '',
            'education': '',
            'summary': '',
            'full': text
        }
        
        # Common section headers
        patterns = {
            'skills': r'(?i)(skills?|technical skills?|core competenc\w+)',
            'experience': r'(?i)(experience|work history|employment|professional experience)',
            'education': r'(?i)(education|academic|qualification)',
            'summary': r'(?i)(summary|objective|profile|about)'
        }
        
        # Try to extract sections
        for section, pattern in patterns.items():
            match = re.search(f'{pattern}[:\s]*\n(.*?)(?=\n[A-Z][A-Z\s]+:|$)', text, re.DOTALL)
            if match:
                sections[section] = match.group(1).strip()
        
        return sections
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\-\+\#\@\.\,\:\;\(\)\/]', '', text)
        return text.strip()